import base64

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from io import BytesIO


def generate_word_report(drawing_id, filename, check_result, image_base64, created_at):
    # Создаем новый документ Word
    doc = Document()

    # Заголовок отчета (кириллица работает идеально!)
    title = doc.add_heading('Отчет по проверке чертежа', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем информацию о чертеже
    doc.add_heading('Информация о чертеже', level=1)

    # Создаем таблицу для информации
    info_table = doc.add_table(rows=4, cols=2)

    # Заполняем таблицу
    info_table.cell(0, 0).text = 'ID чертежа:'
    info_table.cell(0, 1).text = str(drawing_id)

    info_table.cell(1, 0).text = 'Название файла:'
    info_table.cell(1, 1).text = filename

    info_table.cell(2, 0).text = 'Дата проверки:'
    info_table.cell(2, 1).text = created_at

    info_table.cell(3, 0).text = 'Статус:'
    info_table.cell(3, 1).text = 'Проверен'

    # Добавляем раздел с результатами проверки
    doc.add_heading('Результат проверки', level=1)

    # Добавляем текст результата с сохранением переносов строк
    result_paragraph = doc.add_paragraph()
    lines = check_result.split('\n')

    for i, line in enumerate(lines):
        if i > 0:
            result_paragraph.add_run().add_break()  # Добавляем перенос строки между абзацами
        result_paragraph.add_run(line)

    # Добавляем обработанное изображение если есть
    if image_base64:
        try:
            doc.add_heading('Обработанное изображение', level=1)

            # Декодируем base64 и сохраняем во временный файл
            image_data = base64.b64decode(image_base64)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_data)
                temp_file_path = temp_file.name

            # Добавляем изображение в документ
            doc.add_picture(temp_file_path, width=Inches(5.0))

            # Удаляем временный файл
            os.unlink(temp_file_path)

            # Добавляем подпись к изображению
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            # В случае ошибки добавляем сообщение
            error_para = doc.add_paragraph()
            error_para.add_run(f'Ошибка при добавлении изображения: {str(e)}')

    # Сохраняем документ в buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer