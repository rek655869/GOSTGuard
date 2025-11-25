from flask import Flask, request, jsonify
from PIL import Image, ImageDraw
import io
import numpy as np
import base64
import uuid
from flask_cors import CORS
from process_image import process_image
from process_arrow_heads import process_arrow_heads
from process_arrow_distances import process_arrow_distances
from process_text import process_text
from io import BytesIO
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from ultralytics import YOLO
import math

app = Flask(__name__)
CORS(app)

# Загружаем модель один раз при старте
model = YOLO("best.pt")


def get_image_from_request(file):
    """Универсальная функция для получения изображения из запроса"""
    try:
        # Сохраняем файл в памяти для многократного использования
        file_bytes = file.read()
        file_stream = io.BytesIO(file_bytes)
        image = Image.open(file_stream).convert("RGB")
        image_np = np.array(image)
        return image, image_np, file_stream
    except Exception as e:
        raise Exception(f"Ошибка чтения изображения: {str(e)}")


def create_final_image_with_all_annotations(original_image, processed_image,
                                            arrow_heads_violations_data,
                                            arrow_distances_violations_data,
                                            text_violations_data):
    """
    Создает финальное изображение со всеми аннотациями разных типов
    """
    # Начинаем с оригинального изображения
    final_image = original_image.copy()
    draw = ImageDraw.Draw(final_image)

    # Разные цвета для разных типов нарушений
    colors = {
        'arrow_heads': 'pink',
        'arrow_distances': 'blue',
        'text': 'green',
        'frame': 'orange'
    }

    # 1. Рисуем нарушения наконечников стрелок (красные прямоугольники)
    if arrow_heads_violations_data and len(arrow_heads_violations_data) > 0:
        results = model.predict(np.array(original_image), imgsz=640)
        boxes = results[0].boxes.xyxy.cpu().numpy()
        classes = results[0].boxes.cls.cpu().numpy()

        arrows = [box for box, cls in zip(boxes, classes) if cls == 0]

        for i, arrow in enumerate(arrows):
            draw.rectangle([arrow[0], arrow[1], arrow[2], arrow[3]],
                           outline=colors['arrow_heads'], width=3)
            # Подпись для наконечника
            draw.text((arrow[0], arrow[1] - 20), f"Strelka {i + 1}",
                      fill=colors['arrow_heads'])

    # 2. Рисуем нарушения расстояний (синие линии)
    if arrow_distances_violations_data and len(arrow_distances_violations_data) > 0:
        results = model.predict(np.array(original_image), imgsz=640)
        boxes = results[0].boxes.xyxy.cpu().numpy()
        classes = results[0].boxes.cls.cpu().numpy()

        arrows = [box for box, cls in zip(boxes, classes) if cls == 0]
        objects = [box for box, cls in zip(boxes, classes) if cls == 1]

        for i, arrow in enumerate(arrows):
            arrow_center = [(arrow[0] + arrow[2]) / 2, (arrow[1] + arrow[3]) / 2]
            min_distance_px = float('inf')
            closest_obj = None

            for obj in objects:
                obj_center = [(obj[0] + obj[2]) / 2, (obj[1] + obj[3]) / 2]
                distance_px = math.dist(arrow_center, obj_center)
                if distance_px < min_distance_px:
                    min_distance_px = distance_px
                    closest_obj = obj

    # 3. Рисуем нарушения текста (зеленые прямоугольники)
    if text_violations_data and len(text_violations_data) > 0:
        results = model.predict(np.array(original_image), imgsz=640)
        boxes = results[0].boxes.xyxy.cpu().numpy()
        classes = results[0].boxes.cls.cpu().numpy()

        texts = [box for box, cls in zip(boxes, classes) if cls == 2]

        for i, text in enumerate(texts):
            draw.rectangle([text[0], text[1], text[2], text[3]],
                           outline=colors['text'], width=3)
            # Подпись для текста
            draw.text((text[0], text[1] - 20), f"Text {i + 1}",
                      fill=colors['text'])



    return final_image


@app.route('/upload', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    session_id = request.form.get('session_id', str(uuid.uuid4()))

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        # Читаем изображение один раз
        original_image, image_np, file_stream = get_image_from_request(file)

        # 1. ПРОВЕРКА РАМКИ
        processed_image, number, frame_text = process_image(original_image)

        # 2. ПРОВЕРКА НАКОНЕЧНИКОВ СТРЕЛОК
        arrow_heads_violations, arrow_heads_stats, arrow_heads_text, _ = process_arrow_heads(image_np, model)

        # 3. ПРОВЕРКА РАССТОЯНИЙ
        arrow_distances_violations, arrow_distances_stats, arrow_distances_text, _ = process_arrow_distances(image_np,
                                                                                                             model)

        # 4. ПРОВЕРКА ТЕКСТА
        text_violations, text_warnings, text_stats, text_text, _ = process_text(image_np, model)

        # финальное изображение все со всем
        final_image = create_final_image_with_all_annotations(
            original_image=original_image,
            processed_image=processed_image,
            arrow_heads_violations_data=arrow_heads_violations,
            arrow_distances_violations_data=arrow_distances_violations,
            text_violations_data=text_violations
        )

        # ОБЪЕДИНЯЕМ ВСЕ РЕЗУЛЬТАТЫ
        all_violations = (arrow_heads_violations + arrow_distances_violations + text_violations)

        # Формируем общий текст результата
        combined_text = f"""📐 ПРОВЕРКА РАМКИ:
{frame_text}

🎯 ПРОВЕРКА НАКОНЕЧНИКОВ СТРЕЛОК (ГОСТ 2.307-68):
{arrow_heads_text}

📏 ПРОВЕРКА РАССТОЯНИЙ (ГОСТ 2.307-68):
{arrow_distances_text}

📝 ПРОВЕРКА ТЕКСТА (ГОСТ 2.304-81):
{text_text}"""

        # Конвертируем финальное изображение
        buffered = io.BytesIO()
        final_image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')

        return jsonify({
            'success': True,
            'image_base64': img_str,
            'text': combined_text,
            'number': number,
            'session_id': session_id,
            'full_report': {
                'frame_check': {
                    'result': frame_text,
                    'number': number
                },
                'arrow_heads_check': {
                    'result': arrow_heads_text,
                    'violations': arrow_heads_violations,
                    'statistics': arrow_heads_stats,
                    'gost_standard': 'ГОСТ 2.307-68'
                },
                'arrow_distances_check': {
                    'result': arrow_distances_text,
                    'violations': arrow_distances_violations,
                    'statistics': arrow_distances_stats,
                    'gost_standard': 'ГОСТ 2.307-68'
                },
                'text_check': {
                    'result': text_text,
                    'violations': text_violations,
                    'warnings': text_warnings,
                    'statistics': text_stats,
                    'gost_standard': 'ГОСТ 2.304-81'
                },
                'summary': {
                    'total_violations': len(all_violations),
                    'has_violations': len(all_violations) > 0
                }
            }
        })

    except Exception as e:
        print(f"Ошибка в /upload: {str(e)}")
        import traceback
        print(f"Трассировка: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


@app.route('/generate_report', methods=['POST'])
def generate_report():
    try:
        data = request.get_json()

        drawing_id = data.get('drawing_id')
        filename = data.get('filename')
        check_result = data.get('check_result')
        image_base64 = data.get('image_base64')
        created_at = data.get('created_at')

        if not all([drawing_id, filename, check_result, image_base64]):
            return jsonify({'error': 'Missing required data'}), 400

        # Генерируем Word отчет
        doc_buffer = generate_word_report(
            drawing_id=drawing_id,
            filename=filename,
            check_result=check_result,
            image_base64=image_base64,
            created_at=created_at
        )

        # Кодируем Word в base64 для отправки
        doc_base64 = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')

        return jsonify({
            'success': True,
            'doc_base64': doc_base64,
            'filename': f'report_{drawing_id}.docx'
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def generate_word_report(drawing_id, filename, check_result, image_base64, created_at):
    # Создаем новый документ Word
    doc = Document()

    # Заголовок отчета (кириллица работает идеально!)
    title = doc.add_heading('Отчет по проверке чертежа', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем информацию о чертеже
    doc.add_heading('Информация о чертеже', level=1)

    # Создаем таблицу для информации
    info_table = doc.add_table(rows=4, cols=2)

    # Заполняем таблицу
    info_table.cell(0, 0).text = 'ID чертежа:'
    info_table.cell(0, 1).text = str(drawing_id)

    info_table.cell(1, 0).text = 'Название файла:'
    info_table.cell(1, 1).text = filename

    info_table.cell(2, 0).text = 'Дата проверки:'
    info_table.cell(2, 1).text = created_at

    info_table.cell(3, 0).text = 'Статус:'
    info_table.cell(3, 1).text = 'Проверен'

    # Добавляем раздел с результатами проверки
    doc.add_heading('Результат проверки', level=1)

    # Добавляем текст результата с сохранением переносов строк
    result_paragraph = doc.add_paragraph()
    lines = check_result.split('\n')

    for i, line in enumerate(lines):
        if i > 0:
            result_paragraph.add_run().add_break()  # Добавляем перенос строки между абзацами
        result_paragraph.add_run(line)

    # Добавляем обработанное изображение если есть
    if image_base64:
        try:
            doc.add_heading('Обработанное изображение', level=1)

            # Декодируем base64 и сохраняем во временный файл
            image_data = base64.b64decode(image_base64)

            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                temp_file.write(image_data)
                temp_file_path = temp_file.name

            # Добавляем изображение в документ
            doc.add_picture(temp_file_path, width=Inches(5.0))

            # Удаляем временный файл
            os.unlink(temp_file_path)

            # Добавляем подпись к изображению
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            # В случае ошибки добавляем сообщение
            error_para = doc.add_paragraph()
            error_para.add_run(f'Ошибка при добавлении изображения: {str(e)}')

    # Сохраняем документ в buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)